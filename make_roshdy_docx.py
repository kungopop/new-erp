from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import re

md_path = Path('/home/user/ROSHDY_FULL_PROJECT_DOCUMENTATION.md')
out_path = Path('/home/user/ROSHDY_FULL_PROJECT_DOCUMENTATION.docx')
text = md_path.read_text(encoding='utf-8')

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    bidi.set(qn('w:val'), '1')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_cell_rtl(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    tcPr.append(bidi)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        set_rtl(p)


def add_para(doc, txt='', style=None, bold=False):
    p = doc.add_paragraph(style=style)
    set_rtl(p)
    r = p.add_run(txt)
    r.font.name = 'Arial'
    r._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
    r.font.size = Pt(11)
    r.bold = bold
    return p


def clean_inline(s):
    s = s.replace('**','').replace('`','')
    # keep markdown links as text label (url)
    s = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', s)
    return s


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.6)
sec.bottom_margin = Inches(0.6)
sec.left_margin = Inches(0.55)
sec.right_margin = Inches(0.55)

# Normal style
style = doc.styles['Normal']
style.font.name = 'Arial'
style._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
style.font.size = Pt(11)

lines = text.splitlines()
i=0
while i < len(lines):
    line = lines[i]
    if not line.strip():
        i += 1
        continue
    if line.startswith('```'):
        code=[]; i+=1
        while i < len(lines) and not lines[i].startswith('```'):
            code.append(lines[i]); i+=1
        p=add_para(doc, '\n'.join(code))
        for run in p.runs:
            run.font.name='Consolas'; run._element.rPr.rFonts.set(qn('w:cs'), 'Consolas'); run.font.size=Pt(9)
        i+=1
        continue
    if line.startswith('# '):
        p=doc.add_heading(clean_inline(line[2:].strip()), level=0); set_rtl(p)
        i+=1; continue
    if line.startswith('## '):
        p=doc.add_heading(clean_inline(line[3:].strip()), level=1); set_rtl(p)
        i+=1; continue
    if line.startswith('### '):
        p=doc.add_heading(clean_inline(line[4:].strip()), level=2); set_rtl(p)
        i+=1; continue
    if line.startswith('#### '):
        p=doc.add_heading(clean_inline(line[5:].strip()), level=3); set_rtl(p)
        i+=1; continue
    if line.strip() == '---':
        doc.add_paragraph('')
        i+=1; continue
    # Table
    if line.strip().startswith('|') and i+1 < len(lines) and lines[i+1].strip().startswith('|') and '---' in lines[i+1]:
        headers=[clean_inline(x.strip()) for x in line.strip().strip('|').split('|')]
        i += 2
        rows=[]
        while i < len(lines) and lines[i].strip().startswith('|'):
            rows.append([clean_inline(x.strip()) for x in lines[i].strip().strip('|').split('|')])
            i+=1
        table=doc.add_table(rows=1, cols=len(headers))
        table.alignment=WD_TABLE_ALIGNMENT.CENTER
        table.style='Table Grid'
        for c,h in zip(table.rows[0].cells, headers):
            c.text=h
            set_cell_rtl(c)
            for p in c.paragraphs:
                for r in p.runs: r.bold=True
        for row in rows:
            cells=table.add_row().cells
            for c,val in zip(cells,row):
                c.text=val
                set_cell_rtl(c)
        continue
    # Bullet
    if line.lstrip().startswith('- '):
        p=doc.add_paragraph(style='List Bullet')
        set_rtl(p)
        r=p.add_run(clean_inline(line.lstrip()[2:].strip()))
        r.font.name='Arial'; r._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
        i+=1; continue
    # Numbered-ish
    if re.match(r'^\d+\.\s+', line.strip()):
        p=doc.add_paragraph(style='List Number')
        set_rtl(p)
        r=p.add_run(clean_inline(re.sub(r'^\d+\.\s+','',line.strip())))
        r.font.name='Arial'; r._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
        i+=1; continue
    add_para(doc, clean_inline(line.strip()))
    i+=1

# Add final note
p=add_para(doc, 'تم إعداد هذا المستند للاستخدام الإداري والفني، ويمكن تحديثه بعد كل مرحلة تطوير جديدة.', bold=True)

doc.save(out_path)
print(out_path)
